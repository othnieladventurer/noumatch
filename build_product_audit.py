"""
NouMatch Product & Originality Audit — June 2026
Run: python build_product_audit.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\Othniel\Claude\Projects\Noumatch Executive Office\NouMatch_Product_Audit_June2026.docx"

WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
NAVY   = RGBColor(0x1A, 0x1A, 0x2E)
GRAY   = RGBColor(0x55, 0x55, 0x55)
RED    = RGBColor(0xD8, 0x2B, 0x2B)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def h(doc, text, level=1, color=NAVY):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
    return p

def para(doc, text, bold=False, size=10, color=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p

def metric_row(doc, items):
    # items = list of (label, value, bg_hex)
    t = doc.add_table(rows=2, cols=len(items))
    t.style = 'Table Grid'
    for i, (label, value, bg) in enumerate(items):
        vc = t.cell(0, i)
        lc = t.cell(1, i)
        set_cell_bg(vc, bg)
        set_cell_bg(lc, '2C2C3E')
        vp = vc.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr = vp.add_run(value)
        vr.bold = True; vr.font.size = Pt(20); vr.font.color.rgb = WHITE
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run(label)
        lr.font.size = Pt(8); lr.font.color.rgb = RGBColor(0xAA,0xAA,0xCC)
    doc.add_paragraph()

def table2(doc, rows, h1, h2, w1=3.0, w2=4.5):
    t = doc.add_table(rows=1+len(rows), cols=2)
    t.style = 'Table Grid'
    for i, hdr in enumerate([h1, h2]):
        c = t.rows[0].cells[i]
        set_cell_bg(c, '1A1A2E')
        r = c.paragraphs[0].add_run(hdr)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE
    for ri, (c1, c2) in enumerate(rows):
        row = t.rows[ri+1]
        bg = 'F5F5F5' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate([c1, c2]):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            r = cell.paragraphs[0].add_run(str(val))
            r.font.size = Pt(9)
        row.cells[0].width = Inches(w1)
        row.cells[1].width = Inches(w2)
    doc.add_paragraph()

def issue(doc, sev, title, finding, rec):
    colors = {'CRITICAL':'D82B2B','HIGH':'E67E22','MEDIUM':'E6B800','LOW':'1E8A44'}
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    tc = t.cell(0,0); tc.width = Inches(1.2)
    set_cell_bg(tc, colors.get(sev,'888888'))
    tp = tc.paragraphs[0]; tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(sev); tr.bold=True; tr.font.size=Pt(8); tr.font.color.rgb=WHITE
    tp2 = tc.add_paragraph(title); tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp2.runs[0].bold=True; tp2.runs[0].font.size=Pt(9); tp2.runs[0].font.color.rgb=WHITE
    bc = t.cell(0,1); bc.width = Inches(5.8)
    set_cell_bg(bc, 'FAFAFA')
    bp = bc.paragraphs[0]
    bf = bp.add_run('Finding: '); bf.bold=True; bf.font.size=Pt(9)
    bp.add_run(finding).font.size=Pt(9)
    bp2 = bc.add_paragraph()
    br = bp2.add_run('Recommendation: '); br.bold=True; br.font.size=Pt(9)
    bp2.add_run(rec).font.size=Pt(9)
    doc.add_paragraph()

# ── Build document ───────────────────────────────────────────
doc = Document()
s = doc.sections[0]
s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)

# Cover
cp = doc.add_paragraph()
cp.paragraph_format.space_before = Pt(36)
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = cp.add_run('NOUMATCH EXECUTIVE OFFICE')
cr.font.size = Pt(10); cr.bold = True; cr.font.color.rgb = GRAY

tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr2 = tp.add_run('Product & Originality Audit')
tr2.font.size = Pt(26); tr2.bold = True; tr2.font.color.rgb = NAVY

dp = doc.add_paragraph()
dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
dp.add_run('June 14, 2026  ·  Confidential').font.color.rgb = GRAY
doc.add_paragraph()

# ── 1. Executive Summary ─────────────────────────────────────
h(doc, '1. Executive Summary')
para(doc,
    'This audit covers the NouMatch product as of June 14, 2026, based on live production '
    'data from the admin dashboard and full screenshot analysis of the user-facing mobile app. '
    'NouMatch has 85 registered users, 955 total swipes, at least 1 confirmed active match, '
    'and real conversations happening. The platform works. However, the marketplace is in a '
    'pre-liquidity state: DAU is near zero, 46 users have never matched, and the discovery '
    'pool is exhausted for most of the base.', after=8)
para(doc,
    'On originality: the landing page, branding, registration, and profile depth are '
    'meaningfully differentiated from Tinder and Bumble. The core discovery mechanic — '
    'full-screen swipe card with X/heart buttons — is a direct copy of Tinder\'s signature UX. '
    'This is the highest-priority design risk. This report identifies what to protect, '
    'what to change, and in what order.', after=14)

metric_row(doc, [
    ('TOTAL USERS',       '85',    '1A1A2E'),
    ('ACTIVE TODAY',      '2',     'D82B2B'),
    ('TOTAL SWIPES',      '955',   '2471A3'),
    ('LIKE RATE',         '31.8%', '1E8A44'),
    ('ZERO-MATCH USERS',  '46',    'D82B2B'),
    ('TRUST SCORE',       '100',   '16A085'),
])

# ── 2. Marketplace Health ────────────────────────────────────
h(doc, '2. Marketplace Health')

h(doc, '2.1 User Base', 2)
table2(doc, [
    ('Total registered users',      '85'),
    ('Daily active users (DAU)',     '2 (2.4% DAU rate)'),
    ('Latest signup',               'James Kercivil — June 12, 2026'),
    ('Zero-match users',            '46 (54% of base)'),
    ('Average matches per user',    '0.53'),
    ('Gender balance',              '56.6% women / 43.4% men  ✓ within target'),
    ('Avg profile quality score',   '25.0 / 100'),
    ('Trust health score',          '100.0 — no fraud or abuse incidents'),
    ('Ranking strength',            '93.0'),
    ('Top quality account',         'Othniel Lafond — Score 66'),
], 'Metric', 'Value')

h(doc, '2.2 Swipe Activity', 2)
para(doc,
    'Total swipes: 955. Like rate: 31.8% (304 likes, 651 passes). Activity is extremely '
    'concentrated — Francois Michel accounts for 336 swipes (35% of all platform activity), '
    'Evens Jolibois 188. Peak day was June 9 with 46 swipes, followed by a sharp drop to '
    'near-zero. Most users have swiped fewer than 10 times total. Today: 0 swipes.', after=8)
table2(doc, [
    ('Total swipes all-time',  '955'),
    ('Likes',                  '304 (31.8%)'),
    ('Passes',                 '651 (68.2%)'),
    ('Swipes today',           '0'),
    ('Peak swipe day',         'June 9, 2026 — 46 swipes'),
    ('Avg swipes / day',       '7'),
    ('Top swiper',             'Francois Michel — 336 swipes'),
    ('#2 swiper',              'Evens Jolibois — 188 swipes'),
    ('#3 swiper',              'Only One — 40 swipes'),
], 'Metric', 'Value')

h(doc, '2.3 Conversion Funnel', 2)
para(doc,
    'The funnel collapses completely at the match stage. Of users tracked: 100% viewed '
    'profiles, 33.3% liked, 0% matched, 0% messaged. Likes are not converting to matches '
    'because the mutual-like requirement fails in a small inactive pool. There is at least '
    '1 confirmed active match (Mitchama / Othniel) with real conversation data — the system '
    'works when both users are present. The problem is density, not product.', after=8)
table2(doc, [
    ('Views -> Likes',      '33.3%'),
    ('Likes -> Matches',    '0% (tracked window)'),
    ('Matches -> Messages', '0% (tracked window)'),
    ('Active matches',      'At least 1 confirmed'),
    ('Support tickets',     '9 open'),
    ('Reports filed',       '0'),
    ('Blocks recorded',     '0'),
    ('Top-slot like rate',  '0.0%'),
], 'Funnel Stage', 'Rate')

h(doc, '2.4 Diagnosis', 2)
para(doc,
    'NouMatch is in a classic cold-start liquidity trap. The platform is not broken. '
    'The problem is user density. With 85 users distributed across Haiti, the discovery '
    'pool is exhausted for most members. Users joined, saw everyone available, got no '
    'response, and stopped returning. The fix is geographic concentration — 200 active '
    'users in Port-au-Prince specifically is worth more than 1,000 spread across the '
    'diaspora. The admin Boost/Inject controls are already built for exactly this.', after=16)

# ── 3. Product & UX Audit ───────────────────────────────────
h(doc, '3. Product & UX Audit')

h(doc, '3.1 Landing Page', 2)
para(doc,
    'Strong and original. "Des rencontres vraies. Pour des gens vrais." avoids the '
    'over-promise trap that Tinder and Bumble rely on. Real Port-au-Prince couple photography '
    'in the hero is the right instinct. Testimonials and FAQ are present. One gap: the '
    'homepage renders identically for logged-in and logged-out users, contributing to '
    'DAU collapse — returning users have no reason to engage.', after=8)

h(doc, '3.2 Registration Flow', 2)
para(doc,
    'Well-executed. Split-screen layout with Haitian photography + 3-step form is '
    'aesthetically differentiated. The quote "Commence par te presenter, laisse ta '
    'personnalite briller" sets the right tone. Critical gap: the flow does not enforce '
    'bio completion or photo quality before entering discovery. Users with 0% quality scores '
    'and bathroom selfies enter the feed immediately, degrading everyone\'s experience.', after=8)

h(doc, '3.3 Discovery Feed — Core UX', 2)
para(doc,
    'This is the most critical screen and the most significant originality risk. '
    'Full-screen photo card + X/heart buttons + 5-tab bottom nav is Tinder\'s design, '
    'reproduced precisely. It is functional and familiar, but entirely undifferentiated '
    'and self-defeating at 85 users — card-stack swiping requires a large pool to feel '
    'engaging. At this scale, users exhaust the feed in minutes and do not return.', after=8)

h(doc, '3.4 Profile View', 2)
para(doc,
    'Better than Tinder on depth. Structured sections for bio, gender, verification status, '
    'and action buttons. The tagline "L\'amour vaut le risque" at the bottom of every profile '
    'is the most distinctive and original line in the product — a genuine Haitian cultural '
    'signature. It should be elevated, not buried. Issue: "Non verifie" badge appears on '
    'most profiles with no path to verification, undermining the trust signal.', after=8)

h(doc, '3.5 Own Profile & Settings', 2)
para(doc,
    'Rich structure: photo gallery, bio, height, career, education, passions, music. '
    'This depth is a real differentiator vs Tinder\'s minimal card. Critical bug: '
    '"Not Verified" badge displays in English on a French-language app. All UI strings '
    'must be French. Additional language inconsistencies found in error messages and '
    'notification copy.', after=8)

h(doc, '3.6 Matches & Messaging', 2)
para(doc,
    'The chat interface is clean. NouMatch Support is pinned at the top of the '
    'conversation list — a practical trust feature. Real conversation data confirms '
    'the system works. Minor bug: chat input placeholder "Ecrivez votre me..." is truncated. '
    'The matches screen with circular avatars is sparse but functional.', after=8)

h(doc, '3.7 Admin Dashboard', 2)
para(doc,
    'Genuine competitive asset. The custom admin includes Profile Impressions analytics, '
    'Swipe Stats with daily breakdown, Message Center with support inbox, Reports Workspace, '
    'Ranking Analytics, and User Management with Boost/Reduce/Force Inject visibility controls. '
    'This is more sophisticated than most early-stage dating platforms and gives the team '
    'real operational leverage over the marketplace. Protect and extend this.', after=16)

# ── 4. Originality Assessment ────────────────────────────────
h(doc, '4. Originality Assessment')
para(doc,
    'An honest assessment of what is original, what is borrowed, and what must change '
    'to establish NouMatch as a genuinely distinct platform.', after=12)

h(doc, '4.1 What Is Original', 2)
table2(doc, [
    ('Brand positioning',        '"Des rencontres vraies. Pour des gens vrais." — no other dating app uses this framing.'),
    ('Visual identity',          'Red/navy/cream + Haitian photography. Not Tinder (black/white), not Bumble (yellow).'),
    ('Logo',                     'Two silhouetted Haitian faces inside a heart. Culturally specific, not generic.'),
    ('Landing page',             'Real Haiti couple photography. No American app would use this imagery.'),
    ('Profile tagline',          '"L\'amour vaut le risque" on every profile view. Distinctive and Haitian.'),
    ('Profile depth',            'Career, education, passions, music fields — richer than Tinder by design.'),
    ('Language',                 'French throughout. Not a translation of an English product — a native voice.'),
    ('Admin dashboard',          'Custom-built with marketplace visibility controls. Unique operational capability.'),
    ('NouMatch Support chat',    'Pinned support thread in messaging. Practical and trust-building.'),
], 'Element', 'Assessment', w1=2.2, w2=5.3)

h(doc, '4.2 What Is Borrowed (Must Change)', 2)
table2(doc, [
    ('Full-screen swipe card',   'TINDER. Identical layout: photo fills screen, name/age overlay, X/heart buttons.'),
    ('X/heart action buttons',   'TINDER. Pink heart center button and X pass button are Tinder\'s exact pattern.'),
    ('5-tab bottom navigation',  'TINDER. Decouvrir/Likes/Matches/Bloques/Profil mirrors Tinder\'s tab structure.'),
    ('Circular match avatars',   'TINDER. The matches screen avatar row is the same visual pattern.'),
    ('Notification copy',        'TINDER. "[Name] liked your profile!" is identical to Tinder\'s template language.'),
], 'Element', 'Source & Risk', w1=2.2, w2=5.3)

h(doc, '4.3 Recommended Changes', 2)
table2(doc, [
    ('Replace card-stack with profile feed',
     'Show 4-6 profiles in a scrollable grid or vertical feed. Users tap to open and decide. '
     'Matches how Haitians browse content (Instagram-style). Eliminates Tinder comparison at first glance. '
     'Reduces pool-exhaustion feeling at low user counts.'),
    ('Replace X/heart with French pill buttons',
     '"Interesser" / "Passer" as text buttons at the bottom of the opened profile. '
     'Remove the floating pink heart entirely. Language-native, visually distinct.'),
    ('Simplify nav to 3 tabs + FAB',
     'Home / Messages / Profil as the only tabs. Floating action button for discovery. '
     'Move Blocks into profile settings. Frees a slot for something more engaging.'),
    ('Enforce profile quality gate',
     'Require 1 bio sentence + 2 photos before entering discovery. '
     'Every profile in the feed must have text. Raises quality and breaks the Tinder pattern '
     '(Tinder allows empty profiles; NouMatch should not).'),
    ('Elevate "L\'amour vaut le risque"',
     'Animate as a signature entrance on every profile open. Make it a brand moment, '
     'not a footnote. It is the most original line in the product.'),
    ('Rewrite notification copy in Haitian-French',
     '"Bethssu a aime ton profil" instead of "Bethssu liked your profile!" '
     'The current English-template copy undermines the French brand voice.'),
    ('Open the likes grid',
     'Show who liked you openly, with names visible. '
     'Haitian community values transparency; blurring who liked you feels foreign and creates friction.'),
], 'Change', 'Rationale', w1=2.5, w2=5.0)

# ── 5. Issues Register ───────────────────────────────────────
h(doc, '5. Issues Register')

issue(doc, 'CRITICAL', 'Pool Exhausted',
    '85 users, 46 with zero matches, DAU near zero. Most users have seen the entire feed. '
    'Swipe activity peaked June 9 and collapsed to zero since.',
    'Geographic concentration campaign: target 150 additional Port-au-Prince users in 30 days. '
    'Use admin Boost controls to re-inject high-quality inactive profiles immediately.')

issue(doc, 'CRITICAL', 'Discovery = Tinder Clone',
    'Full-screen swipe card + X/heart buttons + 5-tab nav is Tinder\'s design, reproduced exactly. '
    'First-time users recognize it immediately.',
    'Redesign discovery as a profile feed (grid or vertical scroll). '
    'Replace action buttons with French-language pill buttons. Implement in next product sprint.')

issue(doc, 'HIGH', 'Profile Quality Too Low',
    '46% average profile completion. Multiple users in the discovery feed have no bio '
    'and one low-quality photo. This directly suppresses like rates.',
    'Block incomplete profiles from discovery. Use admin "Request bio/photo update" controls '
    'on existing profiles. Send REMINDER notifications to inactive users.')

issue(doc, 'HIGH', 'Logged-In Home = Marketing Page',
    'After login, "/" returns the public homepage. No authenticated dashboard exists. '
    'Users who log in from the homepage see the same page they just left.',
    'Create a logged-in home state showing match activity, new likes, and a feed re-entry CTA. '
    'This is a direct contributor to DAU collapse.')

issue(doc, 'HIGH', 'Language Inconsistencies',
    '"Not Verified" in English on a French app. "Invalide email ou password" is grammatically '
    'wrong French. Chat input text is truncated. Notification copy uses English templates.',
    'Audit all UI strings. All user-facing text must be French. '
    'Priority: verification badge, error messages, notifications, onboarding.')

issue(doc, 'MEDIUM', 'Verification Not Driven',
    'Majority of users are unverified. The badge exists but is not required or incentivized. '
    'Unverified profiles reduce trust for everyone in the feed.',
    'Add verification prompt in onboarding. Give verified users a visible ranking boost. '
    'Surface "Profils verifies" as a platform promise.')

issue(doc, 'MEDIUM', '9 Open Support Tickets',
    'Message Center shows 9 tickets, some open since May 17. '
    '"Staff replied" status on tickets with no actual message content suggests a logging bug.',
    'Review and close all 9 tickets this week. Establish 24h response SLA. Fix logging bug.')

# ── 6. 90-Day Roadmap ───────────────────────────────────────
h(doc, '6. 90-Day Roadmap')
table2(doc, [
    ('Days 1-7',
     'Close all 9 support tickets. Fix all French language errors. '
     'Boost top 20 quality profiles via admin controls. '
     'Send REMINDER notifications to all users inactive >14 days.'),
    ('Days 8-30',
     'Launch Port-au-Prince geographic concentration campaign (+150 local users). '
     'Enforce profile quality gate before discovery entry. '
     'Fix logged-in home state. Rewrite notification copy in Haitian-French.'),
    ('Days 31-60',
     'Redesign discovery feed: replace card-stack with profile grid. '
     'Replace X/heart buttons with French pill actions. '
     'Simplify nav to 3 tabs. Drive verification with onboarding prompt and ranking incentive.'),
    ('Days 61-90',
     'Target 300 active Port-au-Prince users. Re-evaluate match rate and DAU. '
     'If like rate >15%, activate first paid visibility boost feature. '
     'If match rate >5%, begin diaspora expansion (Miami, Montreal, NYC).'),
], 'Phase', 'Actions', w1=1.2, w2=6.3)

# ── 7. Strengths ─────────────────────────────────────────────
h(doc, '7. Confirmed Strengths')
table2(doc, [
    ('Gender balance',       '56.6% women / 43.4% men. Within target. Rare for early-stage dating platforms.'),
    ('Zero abuse incidents', 'Trust score 100.0. No reports, no blocks. Clean community at launch.'),
    ('Real activity',        '955 swipes, confirmed matches, active conversations. The product works.'),
    ('Admin dashboard',      'Custom-built with visibility controls, impression analytics, support inbox.'),
    ('Brand identity',       'Haitian photography, French voice, red/navy palette. Genuinely original.'),
    ('Profile richness',     'Career, education, passions, music fields create more depth than Tinder.'),
    ('Support chat feature', 'NouMatch Support pinned in messaging. Practical trust-building.'),
    ('Landing page copy',    '"Des rencontres vraies. Pour des gens vrais." Strong and original.'),
    ('"L\'amour vaut le risque"', 'The most distinctive line in the product. A genuine cultural signature.'),
], 'Strength', 'Evidence')

doc.add_paragraph()
fp = doc.add_paragraph()
fp.paragraph_format.space_before = Pt(12)
fr = fp.add_run('Prepared by: Noumatch Executive Office (NEO)  |  June 14, 2026  |  Confidential')
fr.font.size = Pt(8)
fr.font.color.rgb = GRAY
fr.italic = True

doc.save(OUT)
print(f"Done. Saved to: {OUT}")
